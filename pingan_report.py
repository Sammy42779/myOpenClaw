from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('平安产险研究报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加日期
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_run = date_para.add_run(f'报告日期：{datetime.now().strftime("%Y年%m月%d日")}')
date_run.font.size = Pt(10)
date_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# 一、公司概况
doc.add_heading('一、公司概况', 1)
doc.add_paragraph(
    '中国平安财产保险股份有限公司（简称"平安产险"）于1988年诞生于深圳蛇口，是中国平安保险（集团）股份有限公司旗下成员，'
    '也是中国金融保险业首家引入外资的企业。作为中国平安集团实现多元化发展的起点，平安产险历经37年发展，已成为中国财产保险行业的领军企业。'
)

# 二、市场地位
doc.add_heading('二、市场地位', 1)
doc.add_paragraph(
    '• 中国第二大财产保险公司（以原保险保费收入衡量）\n'
    '• 2024年实现原保险保费收入3218.2亿元，同比增长6.5%\n'
    '• 连续十四年荣膺中国品牌力指数车险第一品牌\n'
    '• 在中国银保信颁布的车险服务质量指数评价中位列财产险行业前列'
)

# 三、业务范围
doc.add_heading('三、业务范围', 1)
doc.add_paragraph(
    '截至目前，平安产险开发、经营的主险已超过1000种，经营业务范围涵盖：'
)
business_list = [
    '机动车辆保险（车险）',
    '企业财产保险（企财险）',
    '家庭财产保险（家财险）',
    '工程保险',
    '货物运输保险',
    '责任保险',
    '信用保险',
    '保证保险',
    '农业保险（农险）',
    '健康保险',
    '意外伤害保险',
    '国际再保险业务'
]
for item in business_list:
    doc.add_paragraph(item, style='List Bullet')

# 四、金融五篇大文章
doc.add_heading('四、金融"五篇大文章"（2024年成果）', 1)

# 科技金融
doc.add_heading('1. 科技金融', 2)
doc.add_paragraph(
    '累计为电子信息、生物与医药、航空航天、新材料、高新技术、新能源、资源及环境、先进制造等重点领域的'
    '6.9万家科技型企业客户提供超9万亿元的风险保障，助力新质生产力发展。'
)

# 绿色金融
doc.add_heading('2. 绿色金融', 2)
doc.add_paragraph(
    '• 已为2万余棵古树名木提供风险保障超3亿元\n'
    '• 自研的"鹰眼系统"灾害管理平台发送灾害预警信息超105亿次，覆盖超6700万个人及企业客户\n'
    '• 联合中国灾害防御协会推出首本《城市重点消防安全与火灾防控指引》系列科普丛书'
)

# 普惠金融
doc.add_heading('3. 普惠金融', 2)
doc.add_paragraph(
    '• 累计开发超6500款小微保险产品，覆盖逾20个行业\n'
    '• 为240万余家小微企业提供风险保障超220万亿元\n'
    '• 在全国落地563个平安振兴保项目，撬动产业发展资金111.7亿元，带动约150万农户增收144.7亿元\n'
    '• 组织100场乡村振兴助农直播协销活动，协销总金额超2.5亿元，较2023年增长1041%'
)

# 养老金融
doc.add_heading('4. 养老金融', 2)
doc.add_paragraph(
    '已在8个省18个区县承办长期护理保险项目，服务参保群众超3000万，累计支付护理待遇超2.7亿元。'
)

# 数字金融
doc.add_heading('5. 数字金融', 2)
doc.add_paragraph(
    '深化推进数字经济与实体经济深度融合，持续创新"保险+科技+服务"模式。'
    '2024年10月，平安产险凭借"数智化理赔新模式建设项目"荣获中国人民银行"金融科技发展奖"一等奖且排名第一，'
    '成为行业内首个连续荣获一等奖的保险公司。'
)

# 五、数字化服务
doc.add_heading('五、数字化服务', 1)
doc.add_paragraph(
    '平安产险持续践行"三省"工程（省心、省时、省钱），打造极致理赔服务：'
)
doc.add_paragraph(
    '• 平安好车主App：提供涵盖"车保险、车服务、车主生活"的一站式服务\n'
    '  - 注册用户数超2亿\n'
    '  - 月活跃用户数峰值突破4100万\n'
    '  - 客户好评率超97%\n'
    '• 2024年线上办理赔用户数超1056万\n'
    '• 车险自助理赔最快5分钟结案\n'
    '• 万元以下案件理赔时效同比减少1.4天\n'
    '• 非车理赔客户好评率达96%'
)

# 六、联系方式
doc.add_heading('六、联系方式', 1)
doc.add_paragraph(
    '• 官方网站：https://property.pingan.com\n'
    '• 电话车险服务热线：4008-000-000\n'
    '• 总部地址：广东省深圳市福田区'
)

# 七、总结
doc.add_heading('七、总结', 1)
doc.add_paragraph(
    '平安产险作为中国财产保险行业的领军企业，始终坚持稳定、健康的发展战略，'
    '为个人、家庭、企业提供全方位的保险保障服务。公司在建设社会风险管理和保障体系、'
    '推动实体经济发展、助力新技术、支持创新战略等方面，切实发挥保险业"经济减震器"和"社会稳定器"的功能。'
    '未来，平安产险将继续紧跟国家战略导向，全面提升服务社会经济发展质效，全力做好金融"五篇大文章"，'
    '为谱写中国式现代化的新篇章贡献更大力量。'
)

# 保存文档
output_path = '/Users/sammy/.openclaw/workspace/平安产险研究报告.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
